# libs/core/processor/docx_handler.py
"""
DOCX Handler - DOCX Document Processor

Key Features:
- Metadata extraction (title, author, subject, keywords, created/modified dates, etc.)
- Text extraction (direct parsing via python-docx)
- Table extraction (HTML format preservation, rowspan/colspan support)
- Inline image extraction and local saving
- Chart data extraction (OOXML DrawingML Chart parsing)
- Diagram processing

All processing is done via direct binary parsing through python-docx.
Image OCR is performed in a separate post-processing step.

Fallback Chain:
1. Enhanced DOCX processing (python-docx with BytesIO stream)
2. DOCHandler fallback (for non-ZIP files: RTF, OLE, HTML, etc.)
3. Simple text extraction
4. Error message

Class-based Handler:
- DOCXHandler class inherits from BaseHandler to manage config/image_processor
- Internal methods access via self
"""
import io
import logging
import traceback
import zipfile
from typing import Any, Dict, Optional, Set, TYPE_CHECKING

from docx import Document
from lxml import etree

# Base handler
from libs.core.processor.base_handler import BaseHandler
from libs.core.functions.img_processor import ImageProcessor

if TYPE_CHECKING:
    from libs.core.document_processor import CurrentFile

# docx_helper에서 필요한 것들 import
from libs.core.processor.docx_helper import (
    # Constants
    ElementType,
    # Metadata
    extract_docx_metadata,
    format_metadata,
    # Table
    process_table_element,
    # Paragraph
    process_paragraph_element,
)

logger = logging.getLogger("document-processor")


# ============================================================================
# DOCXHandler Class
# ============================================================================

class DOCXHandler(BaseHandler):
    """
    DOCX Document Processing Handler
    
    Inherits from BaseHandler to manage config and image_processor at instance level.
    
    Fallback Chain:
    1. Enhanced DOCX processing (python-docx with BytesIO stream)
    2. DOCHandler fallback (for non-ZIP files: RTF, OLE, HTML, etc.)
    3. Simple text extraction
    4. Error message
    
    Usage:
        handler = DOCXHandler(config=config, image_processor=image_processor)
        text = handler.extract_text(current_file)
    """
    
    def extract_text(
        self,
        current_file: "CurrentFile",
        extract_metadata: bool = True,
        **kwargs
    ) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            current_file: CurrentFile dict containing file info and binary data
            extract_metadata: Whether to extract metadata
            **kwargs: Additional options
            
        Returns:
            Extracted text (with inline image tags, table HTML)
        """
        file_path = current_file.get("file_path", "unknown")
        self.logger.info(f"DOCX processing: {file_path}")
        
        # Check if file is a valid ZIP (DOCX is a ZIP-based format)
        if self._is_valid_zip(current_file):
            return self._extract_docx_enhanced(current_file, extract_metadata)
        else:
            # Not a valid ZIP, try DOCHandler fallback
            self.logger.warning(f"File is not a valid ZIP, trying DOCHandler fallback: {file_path}")
            return self._extract_with_doc_handler_fallback(current_file, extract_metadata)
    
    def _is_valid_zip(self, current_file: "CurrentFile") -> bool:
        """Check if file is a valid ZIP archive."""
        try:
            file_stream = self.get_file_stream(current_file)
            with zipfile.ZipFile(file_stream, 'r') as zf:
                # Check for DOCX-specific content
                return '[Content_Types].xml' in zf.namelist()
        except (zipfile.BadZipFile, Exception):
            return False
    
    def _extract_with_doc_handler_fallback(
        self,
        current_file: "CurrentFile",
        extract_metadata: bool = True
    ) -> str:
        """
        Fallback to DOCHandler for non-ZIP files.
        
        Handles RTF, OLE, HTML, and other formats that might be
        incorrectly named as .docx files.
        """
        file_path = current_file.get("file_path", "unknown")
        
        try:
            from libs.core.processor.doc_handler import DOCHandler
            
            doc_handler = DOCHandler(
                config=self.config,
                image_processor=self.image_processor
            )
            
            # DOCHandler still uses file_path, so pass it directly
            result = doc_handler.extract_text(current_file, extract_metadata=extract_metadata)
            
            if result and not result.startswith("[DOC"):
                self.logger.info(f"DOCHandler fallback successful for: {file_path}")
                return result
            else:
                # DOCHandler also failed, try simple extraction
                return self._extract_simple_text_fallback(current_file)
                
        except Exception as e:
            self.logger.error(f"DOCHandler fallback failed: {e}")
            return self._extract_simple_text_fallback(current_file)
    
    def _extract_simple_text_fallback(self, current_file: "CurrentFile") -> str:
        """
        Last resort: try to extract any readable text from the file.
        """
        file_path = current_file.get("file_path", "unknown")
        file_data = current_file.get("file_data", b"")
        
        try:
            # Try different encodings
            for encoding in ['utf-8', 'cp949', 'euc-kr', 'latin-1']:
                try:
                    text = file_data.decode(encoding)
                    # Remove binary garbage and control characters
                    import re
                    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
                    text = text.strip()
                    
                    if text and len(text) > 50:  # Must have meaningful content
                        self.logger.info(f"Simple text extraction successful with {encoding}: {file_path}")
                        return text
                except (UnicodeDecodeError, Exception):
                    continue
            
            raise ValueError("Could not decode file with any known encoding")
            
        except Exception as e:
            self.logger.error(f"All extraction methods failed for: {file_path}")
            raise RuntimeError(f"DOCX file processing failed: {file_path}. "
                             f"File is not a valid DOCX, DOC, RTF, or text file.")
    
    def _extract_docx_enhanced(
        self,
        current_file: "CurrentFile",
        extract_metadata: bool = True
    ) -> str:
        """
        Enhanced DOCX processing.
        
        - Document order preservation (body element traversal)
        - Metadata extraction
        - Inline image extraction and local saving
        - Table HTML format preservation (cell merge support)
        - Chart data extraction
        - Page break handling
        """
        file_path = current_file.get("file_path", "unknown")
        self.logger.info(f"Enhanced DOCX processing: {file_path}")

        try:
            # Use BytesIO stream to avoid path encoding issues
            file_stream = self.get_file_stream(current_file)
            doc = Document(file_stream)
            result_parts = []
            processed_images: Set[str] = set()
            current_page = 1
            total_tables = 0
            total_images = 0
            total_charts = 0

            # Metadata extraction
            if extract_metadata:
                metadata = extract_docx_metadata(doc)
                metadata_str = format_metadata(metadata)
                if metadata_str:
                    result_parts.append(metadata_str + "\n\n")
                    self.logger.info(f"DOCX metadata extracted: {list(metadata.keys())}")

            # Start page 1
            page_tag = self.create_page_tag(current_page)
            result_parts.append(f"{page_tag}\n")

            # Traverse body elements in document order
            for body_elem in doc.element.body:
                local_tag = etree.QName(body_elem).localname

                if local_tag == 'p':
                    # Paragraph processing - pass image_processor
                    content, has_page_break, img_count, chart_count = process_paragraph_element(
                        body_elem, doc, processed_images, file_path,
                        image_processor=self.image_processor
                    )

                    if has_page_break:
                        current_page += 1
                        page_tag = self.create_page_tag(current_page)
                        result_parts.append(f"\n{page_tag}\n")

                    if content.strip():
                        result_parts.append(content + "\n")

                    total_images += img_count
                    total_charts += chart_count

                elif local_tag == 'tbl':
                    table_html = process_table_element(body_elem, doc)
                    if table_html:
                        total_tables += 1
                        result_parts.append("\n" + table_html + "\n\n")

                elif local_tag == 'sectPr':
                    continue

            result = "".join(result_parts)
            self.logger.info(f"Enhanced DOCX processing completed: {current_page} pages, "
                           f"{total_tables} tables, {total_images} images, {total_charts} charts")

            return result

        except Exception as e:
            self.logger.error(f"Error in enhanced DOCX processing: {e}")
            self.logger.debug(traceback.format_exc())
            return self._extract_docx_simple_text(current_file)
    
    def _extract_docx_simple_text(self, current_file: "CurrentFile") -> str:
        """Simple text extraction (fallback)."""
        try:
            file_stream = self.get_file_stream(current_file)
            doc = Document(file_stream)
            result_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    result_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        row_texts.append(cell.text.strip())
                    if any(t for t in row_texts):
                        result_parts.append(" | ".join(row_texts))

            return "\n".join(result_parts)

        except Exception as e:
            self.logger.error(f"Error in simple DOCX text extraction: {e}")
            return f"[DOCX file processing failed: {str(e)}]"


__all__ = ["DOCXHandler"]
